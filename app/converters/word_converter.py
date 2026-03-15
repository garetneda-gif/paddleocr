"""Word (DOCX) 导出器 — 使用 bbox 坐标还原版面布局。"""

from __future__ import annotations

from pathlib import Path

from app.converters.base_converter import BaseConverter
from app.converters.layout_analyzer import analyze_page, Paragraph
from app.models import DocumentResult, BlockType


class WordConverter(BaseConverter):
    @property
    def file_extension(self) -> str:
        return ".docx"

    def convert(self, result: DocumentResult, output_path: Path) -> Path:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "SimSun"
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

        for page_idx, page in enumerate(result.pages):
            if page_idx > 0:
                doc.add_page_break()

            if _has_semantic_blocks(page):
                _render_semantic_blocks(doc, page.blocks)
                continue

            paragraphs = analyze_page(page)

            if not paragraphs:
                # 回退：无分析结果时直接输出 blocks
                for block in page.blocks:
                    if block.text:
                        doc.add_paragraph(block.text)
                continue

            has_columns = any(p.column > 0 for p in paragraphs)

            for para in paragraphs:
                if para.block_type == BlockType.TITLE:
                    heading = doc.add_heading(para.text, level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p = doc.add_paragraph()
                    if has_columns and para.column > 0:
                        # 双栏：用缩进模拟
                        if para.column == 2:
                            p.paragraph_format.left_indent = Cm(8)
                    run = p.add_run(para.text)
                    run.font.size = Pt(11)

        # 降级：如果没有 blocks，使用 plain_text
        if not any(page.blocks for page in result.pages) and result.plain_text:
            for line in result.plain_text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

        doc.save(str(output_path))
        return output_path


def _has_semantic_blocks(page) -> bool:
    return any(
        block.block_type != BlockType.PARAGRAPH or block.table_cells
        for block in page.blocks
    )


def _render_semantic_blocks(doc, blocks) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for block in sorted(blocks, key=lambda item: (item.bbox[1], item.bbox[0])):
        if block.table_cells:
            row_count = len(block.table_cells)
            col_count = max((len(row) for row in block.table_cells), default=0)
            if row_count == 0 or col_count == 0:
                continue
            table = doc.add_table(rows=row_count, cols=col_count)
            table.style = "Table Grid"
            for row_idx, row in enumerate(block.table_cells):
                for col_idx, cell in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(cell)
            continue

        text = block.text.strip()
        if not text:
            continue
        if block.block_type == BlockType.TITLE:
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(11)
